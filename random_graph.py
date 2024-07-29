from docx import Document
from graph import Graph
from search import GraphSearch
from random import random, randint
from timeit import repeat
import matplotlib.pyplot as plt

# Define the sizes and probabilities
sizes = [10, 20, 30, 40]
probabilities = [0.2, 0.4, 0.6, 0.8]

# Generate random locations for nodes
node_locations = {}
for n in range(1, 41):
    while True:
        i, j = randint(0, 10), randint(0, 10)
        if (i, j) not in node_locations.values():
            node_locations[n] = (i, j)
            break

graphs = []
for size in sizes:
    for prob in probabilities:
        graph = Graph()
        graph.location = node_locations
        graph.addNodes([graph.createNode(val) for val in list(range(1, size+1))])
        for i in range(1, size+1):
            for j in range(i+1, size+1):
                edge_probability = random()
                if edge_probability < prob:
                    graph.addNeighbour(i, j)
        
        graphs.append(graph)

# Randomly select 10 nodes
randomly_selected_nodes = []
for graph in graphs:
    selected_nodes = []
    if len(graph.adj_list) > 10:
        nodes = list(graph.adj_list.keys())
        for i in range(10):
            while True:
                random_index = randint(0, len(graph.adj_list)-1)
                if nodes[random_index] not in selected_nodes:
                    selected_nodes.append(nodes[random_index])
                    break
    else:
        selected_nodes += list(graph.adj_list.keys())
    randomly_selected_nodes.append(selected_nodes)

# Code snippets for different search algorithms
setup_code = '''import __main__
from graph import Graph
from random import random, randint
from search import GraphSearch'''

search_codes = {
    'DFS': '''
def test():
    GraphSearch.DFS(graph, randomly_selected_nodes[n][i], randomly_selected_nodes[n][j])
''',
    'UCS': '''
def test():
    GraphSearch.UCS(graph, randomly_selected_nodes[n][i], randomly_selected_nodes[n][j])
''',
    'A*': '''
def test():
    GraphSearch.astar(graph, randomly_selected_nodes[n][i], randomly_selected_nodes[n][j])
''',
    'Iterative': '''
def test():
    GraphSearch.iterativeDeepening(graph, randomly_selected_nodes[n][i], randomly_selected_nodes[n][j])
''',
    'BFS': '''
def test():
    GraphSearch.BFS(graph, randomly_selected_nodes[n][i], randomly_selected_nodes[n][j])
''',
    'Greedy': '''
def test():
    GraphSearch.greedySearch(graph, randomly_selected_nodes[n][i], randomly_selected_nodes[n][j])
'''
}

# Measure execution times
all_execution_times = []
for graph_index, graph in enumerate(graphs):
    graph_times = []
    for _ in range(10):
        for i in range(10):
            for j in range(i+1, 10):
                time_taken = sum(repeat(setup=setup_code, stmt=search_codes['DFS'], repeat=5, number=1, globals=globals())) / 5
                graph_times.append(time_taken)
    all_execution_times.append(graph_times)

# Path lengths for different algorithms
all_path_lengths = []
for graph_index, graph in enumerate(graphs):
    dfs_paths = 0
    a_star_paths = 0
    iterative_paths = 0
    greedy_paths = 0
    for _ in range(10):
        for i in range(10):
            for j in range(i+1, 10):
                dfs_paths += len(GraphSearch.DFS(graph, randomly_selected_nodes[graph_index][i], randomly_selected_nodes[graph_index][j])) - 1
                a_star_paths += len(GraphSearch.astar(graph, randomly_selected_nodes[graph_index][i], randomly_selected_nodes[graph_index][j])) - 1
                iterative_paths += len(GraphSearch.iterativeDeepening(graph, randomly_selected_nodes[graph_index][i], randomly_selected_nodes[graph_index][j], 5)) - 1
                greedy_paths += len(GraphSearch.greedySearch(graph, randomly_selected_nodes[graph_index][i], randomly_selected_nodes[graph_index][j])) - 1
    all_path_lengths.append([dfs_paths, a_star_paths, a_star_paths, iterative_paths, a_star_paths, greedy_paths])

# Create a Word document and record the time and path length
document = Document()
# Record execution times
table = document.add_table(rows=1, cols=7)
hdr_cells = table.rows[0].cells

algorithm_names = ['DFS', 'UCS', 'A*', 'Iterative', 'BFS', 'Greedy']

hdr_cells[0].text = ''
for i in range(len(algorithm_names)):
    hdr_cells[i+1].text = algorithm_names[i]

for i in range(16):
    row_cells = table.add_row().cells
    row_cells[0].text = 'Graph ' + str(i+1)
    for j in range(6):
        row_cells[j+1].text = str(all_execution_times[i][j])





document.save('random_graph.docx')

# Test each algorithm and record the results
results = []

for code in search_codes.values():
    result = repeat(setup=setup_code, stmt=code, repeat=5, number=1)
    results.append(result)

# Record the average time taken for each algorithm and graph
averages = [sum(result) / len(result) for result in results]

# Write results to the Word document
document_results = Document()

table = document_results.add_table(rows=1, cols=len(algorithm_names) + 1)
hdr_cells = table.rows[0].cells

hdr_cells[0].text = 'Algorithm'
for i, name in enumerate(algorithm_names):
    hdr_cells[i+1].text = name

for i in range(5):  # Repeat for each experiment
    row_cells = table.add_row().cells
    row_cells[0].text = 'Experiment ' + str(i+1)
    for j in range(len(averages)):
        row_cells[j+1].text = str(results[j][i])

row_cells = table.add_row().cells
row_cells[0].text = 'Average time'
for j in range(len(averages)):
    row_cells[j+1].text = str(averages[j])

document_results.save('average_time_results.docx')

# Plotting execution times
execution_times_array = all_execution_times

# Define colors and algorithm labels
colors = ['red', 'blue', 'green', 'orange', 'purple', 'brown']
algorithm_labels = {0: 'DFS', 1: 'UCS', 2: 'A*', 3: 'Iterative', 4: 'BFS', 5: 'Greedy'}

# Plot time data
fig, ax = plt.subplots()
for i, label in algorithm_labels.items():
    ax.plot(range(1, len(execution_times_array) + 1), [row[i] for row in execution_times_array], marker='o', color=colors[i], label=label)

# Set axis labels and legend
ax.set_xlabel('Number of Graphs')
ax.set_ylabel('Time')
ax.legend(title='Algorithms')

plt.show()

# Plotting for path lengths
path_lengths_array = all_path_lengths

# Plot path data
fig, ax = plt.subplots()
for i, label in algorithm_labels.items():
    ax.plot(range(1, len(path_lengths_array) + 1), [row[i] for row in path_lengths_array], marker='o', color=colors[i], label=label)

# Set axis labels and legend
ax.set_xlabel('Number of Graphs')
ax.set_ylabel('Path Length')
ax.legend(title='Algorithms')

plt.show()
